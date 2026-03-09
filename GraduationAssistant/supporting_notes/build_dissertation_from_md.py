from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(r"C:\Repositories\20260309_DrGraduationAdminAssist")
SOURCE_MD = BASE_DIR / "Dissertation" / "archive_drafts" / "Dissertation_20260309_1448_source.md"
TARGET_DOCX = BASE_DIR / "Dissertation" / "Dissertation_20260309_1452.docx"
BROKEN_DOCX = BASE_DIR / "Dissertation" / "archive_drafts" / "Dissertation_20260309_1452_broken.docx"


def set_run_font(run, name="Batang", size=11, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_styles(doc):
    for style_name, size in [
        ("Normal", 11),
        ("Heading 1", 14),
        ("Heading 2", 12),
        ("Heading 3", 11),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Batang"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Batang")
        style.font.size = Pt(size)


def add_paragraph(doc, text, style="Normal", align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12 if level == 2 else 11, bold=True)
    return p


def add_list(doc, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.4
    run = p.add_run(text)
    set_run_font(run)
    return p


def build():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Missing source markdown: {SOURCE_MD}")

    if TARGET_DOCX.exists():
        TARGET_DOCX.replace(BROKEN_DOCX)

    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    pending_paragraph = []

    def flush_paragraph():
        nonlocal pending_paragraph
        if pending_paragraph:
            text = " ".join(x.strip() for x in pending_paragraph if x.strip()).strip()
            if text:
                add_paragraph(doc, text)
            pending_paragraph = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            add_heading(doc, stripped[2:].strip(), 1)
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            add_heading(doc, stripped[3:].strip(), 2)
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            add_heading(doc, stripped[4:].strip(), 3)
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_list(doc, stripped[2:].strip(), numbered=False)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph()
            text = re.sub(r"^\d+\.\s+", "", stripped)
            add_list(doc, text, numbered=True)
            continue

        pending_paragraph.append(stripped)

    flush_paragraph()
    doc.save(str(TARGET_DOCX))


if __name__ == "__main__":
    build()
