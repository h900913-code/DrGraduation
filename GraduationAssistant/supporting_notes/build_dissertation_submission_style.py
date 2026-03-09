from pathlib import Path
import re
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DIR = Path(r"C:\Repositories\20260309_DrGraduationAdminAssist")
SOURCE_MD = BASE_DIR / "Dissertation" / "archive_drafts" / "Dissertation_20260309_1448_source.md"
TARGET_DOCX = BASE_DIR / "Dissertation" / "Dissertation_20260309_1452.docx"
ARCHIVE_DIR = BASE_DIR / "Dissertation" / "archive_drafts"
ARCHIVE_COPY = ARCHIVE_DIR / "Dissertation_20260309_1452_pre_submission_style.docx"

AUTHOR_KO = "박정현"
AUTHOR_EN = "Jeonghyeon Park"
STUDENT_NO = "2021-37772"
SCHOOL_KO = "서울대학교 환경대학원"
DEPARTMENT_KO = "환경관리학과"
DEGREE_KO = "환경관리학박사학위논문"
EXPECTED_GRAD_MONTH = "2027년 2월"
SUBMISSION_MONTH = "2026년 10월"
APPROVAL_MONTH = "2026년 12월"
ADVISOR = "[지도교수 성명]"
COMMITTEE = ["[위원장]", "[부위원장]", "[위원]", "[위원]", "[위원]"]


def clean_text(text: str) -> str:
    text = text.replace("`", "")
    text = text.replace("\u00a0", " ")
    return text.strip()


def split_sections(text: str):
    sections = []
    current_title = None
    buffer = []
    for line in text.splitlines():
        if line.startswith("## "):
            if current_title is not None:
                sections.append((current_title, "\n".join(buffer).strip()))
            current_title = line[3:].strip()
            buffer = []
        else:
            if current_title is not None:
                buffer.append(line)
    if current_title is not None:
        sections.append((current_title, "\n".join(buffer).strip()))
    return sections


def parse_bullets(text: str):
    items = []
    for line in text.splitlines():
        s = line.strip()
        if s.startswith("- "):
            items.append(clean_text(s[2:]))
    return items


def parse_paragraphs(text: str):
    parts = []
    chunk = []
    for line in text.splitlines():
        s = line.rstrip()
        if not s.strip():
            if chunk:
                parts.append(clean_text(" ".join(x.strip() for x in chunk)))
                chunk = []
            continue
        chunk.append(s)
    if chunk:
        parts.append(clean_text(" ".join(x.strip() for x in chunk)))
    return parts


def parse_chapter_body(text: str):
    blocks = []
    current_subtitle = None
    current_lines = []

    def flush():
        nonlocal current_lines, current_subtitle
        if current_subtitle is not None or current_lines:
            blocks.append((current_subtitle, parse_paragraphs("\n".join(current_lines))))
        current_subtitle = None
        current_lines = []

    for line in text.splitlines():
        if line.startswith("### "):
            flush()
            current_subtitle = clean_text(line[4:])
        else:
            current_lines.append(line)
    flush()
    return blocks


def parse_references(text: str):
    refs = []
    chunk = []
    for line in text.splitlines():
        s = line.strip()
        if not s:
            if chunk:
                refs.append(clean_text(" ".join(chunk)))
                chunk = []
            continue
        chunk.append(s)
    if chunk:
        refs.append(clean_text(" ".join(chunk)))
    return refs


def set_font(run, name="Batang", size=11, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def configure_styles(doc: Document):
    for style_name, size in [("Normal", 11), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Batang"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Batang")
        style.font.size = Pt(size)


def add_paragraph(doc: Document, text: str = "", align=None, size=11, bold=False, italic=False, spacing=1.7):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = spacing
    if align is not None:
        p.alignment = align
    run = p.add_run(clean_text(text))
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc: Document, text: str, level: int = 1, align=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    if align is not None:
        p.alignment = align
    run = p.add_run(clean_text(text))
    set_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_list(doc: Document, text: str, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(clean_text(text))
    set_font(run)
    return p


def add_page_break(doc: Document):
    doc.add_page_break()


def romanize_chapter(title: str):
    mapping = {
        "제1장": "Ⅰ.",
        "제2장": "Ⅱ.",
        "제3장": "Ⅲ.",
        "제4장": "Ⅳ.",
        "제5장": "Ⅴ.",
        "제6장": "Ⅵ.",
        "제7장": "Ⅶ.",
        "제8장": "Ⅷ.",
        "제9장": "Ⅸ.",
    }
    for k, v in mapping.items():
        if title.startswith(k):
            rest = title.replace(k, "", 1).strip()
            return f"{v} {rest}"
    return title


def add_cover(doc: Document, title: str):
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, DEGREE_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    if ":" in title:
        left, right = [x.strip() for x in title.split(":", 1)]
        add_paragraph(doc, left, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, spacing=1.2)
        add_paragraph(doc, f"- {right}", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, spacing=1.2)
    else:
        add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, spacing=1.2)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, EXPECTED_GRAD_MONTH, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, SCHOOL_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, DEPARTMENT_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, AUTHOR_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_page_break(doc)


def add_title_page(doc: Document, title: str):
    if ":" in title:
        left, right = [x.strip() for x in title.split(":", 1)]
        add_paragraph(doc, left, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, spacing=1.2)
        add_paragraph(doc, f"- {right}", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, spacing=1.2)
    else:
        add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, spacing=1.2)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"지도교수 {ADVISOR}", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, f"이 논문을 {DEGREE_KO}으로 제출함", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, SUBMISSION_MONTH, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, SCHOOL_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, DEPARTMENT_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, AUTHOR_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_page_break(doc)


def add_approval_page(doc: Document, title: str):
    if ":" in title:
        left, right = [x.strip() for x in title.split(":", 1)]
        add_paragraph(doc, left, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, spacing=1.2)
        add_paragraph(doc, f"- {right}", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, spacing=1.2)
    else:
        add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, bold=True, spacing=1.2)
    add_paragraph(doc, f"지도교수 {ADVISOR}", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, f"이 논문을 {DEGREE_KO}으로 제출함", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, SUBMISSION_MONTH, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, SCHOOL_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, DEPARTMENT_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, AUTHOR_KO, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"{AUTHOR_KO}의 {DEGREE_KO}을 인준함", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_paragraph(doc, APPROVAL_MONTH, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    for idx, member in enumerate(COMMITTEE):
        label = "위원장" if idx == 0 else "부위원장" if idx == 1 else "위원"
        add_paragraph(doc, f"{label}   {member}   (인)", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    add_page_break(doc)


def add_blank_consent_placeholder(doc: Document):
    add_heading(doc, "원문제공서비스에 대한 동의서", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "최종 제출본 1부에는 도서관 제출용 원문제공서비스 동의서 원본을 합철한다.", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "현재 초안에서는 해당 양식을 포함하지 않고 자리만 표시한다.", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_page_break(doc)


def add_korean_abstract(doc: Document, title: str, text: str, keywords: str):
    add_heading(doc, "국문초록", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, spacing=1.2)
    for paragraph in parse_paragraphs(text):
        add_paragraph(doc, paragraph)
    add_paragraph(doc, f"주요어: {keywords}")
    add_paragraph(doc, f"학번: {STUDENT_NO}")
    add_page_break(doc)


def add_toc(doc: Document, chapters):
    add_heading(doc, "<목차>", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ch in chapters:
        add_paragraph(doc, romanize_chapter(ch), spacing=1.4)
    add_paragraph(doc, "참고문헌", spacing=1.4)
    add_paragraph(doc, "Abstract", spacing=1.4)
    add_page_break(doc)
    add_heading(doc, "표목차", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "현재 초안 단계이므로 실제 표번호 및 페이지는 결과 장 완성 후 자동 정리 예정.")
    add_page_break(doc)
    add_heading(doc, "그림목차", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, "현재 초안 단계이므로 실제 그림번호 및 페이지는 결과 장 완성 후 자동 정리 예정.")
    add_page_break(doc)


def add_chapter(doc: Document, title: str, body_text: str):
    add_heading(doc, romanize_chapter(title), level=1)
    blocks = parse_chapter_body(body_text)
    if blocks and all(subtitle is None for subtitle, _ in blocks):
        for _, paragraphs in blocks:
            for para in paragraphs:
                if para.startswith("[") and para.endswith("]"):
                    add_paragraph(doc, para, italic=True)
                else:
                    add_paragraph(doc, para)
        return

    for subtitle, paragraphs in blocks:
        if subtitle:
            # convert 1. -> 1. and 1) 그대로 유지
            add_heading(doc, subtitle, level=2)
        for para in paragraphs:
            if not para:
                continue
            if para.startswith("[") and para.endswith("]"):
                add_paragraph(doc, para, italic=True)
            elif re.match(r"^\d+\.\s+", para):
                add_list(doc, re.sub(r"^\d+\.\s+", "", para), numbered=True)
            elif para.startswith("- "):
                add_list(doc, para[2:], numbered=False)
            else:
                add_paragraph(doc, para)


def add_references(doc: Document, references):
    add_heading(doc, "참고문헌", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ref in references:
        add_paragraph(doc, ref)


def add_english_abstract(doc: Document, title: str, text: str):
    add_page_break(doc)
    add_heading(doc, "Abstract", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, spacing=1.2)
    add_paragraph(doc, AUTHOR_EN, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"{DEPARTMENT_KO}, {SCHOOL_KO}", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f"Student Number: {STUDENT_NO}", align=WD_ALIGN_PARAGRAPH.CENTER)
    for paragraph in parse_paragraphs(text):
        add_paragraph(doc, paragraph)


def build():
    text = SOURCE_MD.read_text(encoding="utf-8")
    sections = split_sections(text)
    section_map = {title: content for title, content in sections}

    title = clean_text(parse_paragraphs(section_map["제목"])[0])
    draft_notes = parse_bullets(section_map.get("초안 메모", ""))
    korean_abstract = section_map.get("국문초록", "")
    english_abstract = section_map.get("Abstract", "")
    references = parse_references(section_map.get("잠정 참고문헌", ""))

    chapter_titles = [title for title, _ in sections if title.startswith("제")]

    if TARGET_DOCX.exists():
        ARCHIVE_DIR.mkdir(exist_ok=True)
        shutil.copy2(TARGET_DOCX, ARCHIVE_COPY)

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    add_cover(doc, title)
    add_title_page(doc, title)
    add_approval_page(doc, title)
    add_blank_consent_placeholder(doc)
    add_korean_abstract(
        doc,
        title,
        korean_abstract,
        "파리협정체제, 기후변화, 유튜브, 담론 분석, 정서·태도 분석, 네트워크 분석",
    )
    add_toc(doc, chapter_titles)

    # Draft memo page
    add_heading(doc, "초안 메모", level=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    for item in draft_notes:
        add_list(doc, item)
    add_page_break(doc)

    for chapter_title in chapter_titles:
        body = section_map[chapter_title]
        add_chapter(doc, chapter_title, body)

    add_references(doc, references)
    add_english_abstract(doc, title, english_abstract)
    doc.save(str(TARGET_DOCX))


if __name__ == "__main__":
    build()
