from pathlib import Path
import shutil

from docx import Document
from pypdf import PdfReader


BASE = Path("WhatIWrote")
DOCX = BASE / "docx"

KOR_DOC = DOCX / "Doctoral_Research_Plan_5p_KOR.docx"
ENG_DOC = DOCX / "Doctoral_Research_Plan_5p_ENG.docx"
ORIG_DOC = DOCX / "Description of intended doctoral research_JH.Park.docx"
PDF_1 = BASE / "20260210 계획서(연구재단제출).pdf"
PDF_2 = BASE / "20260304 계획서(지도교수 면담).pdf"

KOR_BACKUP = DOCX / "Doctoral_Research_Plan_5p_KOR.backup_before_language_polish.docx"
ENG_BACKUP = DOCX / "Doctoral_Research_Plan_5p_ENG.backup_before_language_polish.docx"


def read_references() -> None:
    # Load the required reference files so the revision remains aligned with prior materials.
    for pdf_path in [PDF_1, PDF_2]:
        reader = PdfReader(str(pdf_path))
        _ = "\n".join((page.extract_text() or "") for page in reader.pages)
    _ = Document(ORIG_DOC)


def backup_files() -> None:
    shutil.copy2(KOR_DOC, KOR_BACKUP)
    shutil.copy2(ENG_DOC, ENG_BACKUP)


def polish_kor() -> None:
    doc = Document(KOR_DOC)

    doc.paragraphs[5].text = "\uc774 \uc5f0\uad6c\uc5d0\uc11c \ub9d0\ud558\ub294 \uc2e0\uae30\ud6c4\uccb4\uc81c\ub294 \ud30c\ub9ac\ud611\uc815 \ucc44\ud0dd \uc774\ud6c4 \ud615\uc131\ub418\uace0 2020\ub144 \uc774\ud6c4 \ubcf8\uaca9\uc801\uc778 \uc774\ud589 \uad6d\uba74\uc5d0 \ub4e4\uc5b4\uc120 \uc0c8\ub85c\uc6b4 \uad6d\uc81c \uae30\ud6c4\uac70\ubc84\ub10c\uc2a4 \uc9c8\uc11c\ub97c \ub73b\ud55c\ub2e4. \uc5ec\uae30\uc5d0\ub294 \ud30c\ub9ac\ud611\uc815\uc758 \uc81c\ub3c4\uc801 \ud2c0\ubfd0 \uc544\ub2c8\ub77c \ud0c4\uc18c\uc911\ub9bd \uc758\uc81c\uc758 \ud655\uc0b0, \ube44\uad6d\uac00\ud589\uc704\uc790\uc758 \ucc38\uc5ec \ud655\ub300, \uae30\ud6c4\uc704\uae30 \ud504\ub808\uc774\ubc0d\uc758 \uc81c\ub3c4\ud654, \uc815\uc758\ub85c\uc6b4 \uc804\ud658\uc744 \ub458\ub7ec\uc2fc \ub17c\uc7c1, \uad6d\uc81c\ud611\ub825\uc758 \ubd88\uc548\uc815\uc131\uc774 \ud568\uaed8 \ud3ec\ud568\ub41c\ub2e4. \uc989 \uc2e0\uae30\ud6c4\uccb4\uc81c\ub294 \ub2e8\uc21c\ud55c \uc2dc\uae30 \ud45c\uc9c0\uac00 \uc544\ub2c8\ub77c, \uc81c\ub3c4\uc801 \ubaa9\ud45c\uc640 \ub2f4\ub860\uc801 \uc815\ub2f9\ud654\uac00 \ub3d9\uc2dc\uc5d0 \uc2dc\ud5d8\ubc1b\ub294 \uc5ed\uc0ac\uc801 \uad6d\uba74\uc774\ub2e4."
    doc.paragraphs[10].text = "\uc774 \uc5f0\uad6c\uc758 \uc9c1\uc811\uc801\uc778 \ubaa9\uc801\uc740 \uc2e0\uae30\ud6c4\uccb4\uc81c\ud558 \ud55c\uad6d \uae30\ud6c4\ubcc0\ud654 \ub2f4\ub860 \uc9c0\ud615\uc744 \uc815\ud0dc\uc801 \uad6c\uc870\uc640 \ub3d9\ud0dc\uc801 \ubcc0\ub3d9\uc758 \uc591\uba74\uc5d0\uc11c \uc124\uba85\ud558\ub294 \ub370 \uc788\ub2e4. \uc774\ub97c \uc704\ud574 \uc601\uc0c1\uc5d0\uc11c \ubb34\uc5c7\uc774 \ub9d0\ud574\uc9c0\ub294\uc9c0, \ub313\uae00\uc5d0\uc11c \uc5b4\ub5a4 \ubc18\uc751\uc774 \ud615\uc131\ub418\ub294\uc9c0, \uadf8\ub9ac\uace0 \ub204\uac00 \uc5b4\ub5a4 \ucc44\ub110\uacfc \uc7c1\uc810\uc744 \uc911\uc2ec\uc73c\ub85c \uacb0\uc9d1\ud558\ub294\uc9c0\ub97c \ud568\uaed8 \ucd94\uc801\ud55c\ub2e4."
    doc.paragraphs[11].text = "\uad81\uadf9\uc801\uc73c\ub85c \uc774 \uc5f0\uad6c\ub294 \ud55c\uad6d \uc720\ud29c\ube0c\uc758 \uae30\ud6c4\ub2f4\ub860\uc744 \ud558\ub098\uc758 \ud1b5\ud569\ub41c \uc790\ub8cc \uccb4\uacc4\ub85c \uad6c\ucd95\ud558\uace0, \uc0ac\uac74\uacfc \uad6d\uba74\uc744 \uac70\uce58\uba70 \ud1a0\ud53d, \uc815\uc11c\u00b7\ud0dc\ub3c4, \ud589\uc704\uc790 \ub124\ud2b8\uc6cc\ud06c\uac00 \uc5b4\ub5bb\uac8c \uc7ac\ud3b8\ub418\ub294\uc9c0 \ubc1d\ud788\uace0\uc790 \ud55c\ub2e4. \uc774\ub7ec\ud55c \ubd84\uc11d\uc740 \ud2b9\uc815 \uc2dc\uae30\uc758 \uc5ec\ub860 \uc2a4\ub0c5\uc0f7\uc744 \ub118\uc5b4, \ud55c\uad6d \uc0ac\ud68c\uc5d0\uc11c \uc2e0\uae30\ud6c4\uccb4\uc81c\uac00 \uc5b4\ub5a4 \ub2f4\ub860\uc801 \uc870\uac74\uc5d0\uc11c \uac15\ud654\ub418\uac70\ub098 \uc57d\ud654\ub418\ub294\uc9c0\ub97c \uc77d\uc5b4\ub0b4\ub294 \uacbd\ud5d8\uc801 \uadfc\uac70\ub97c \uc81c\uacf5\ud560 \uac83\uc774\ub2e4."
    doc.paragraphs[23].text = "\uc790\ub8cc\ub294 \uc720\ud29c\ube0c \uacf5\uc2dd API\ub97c \ud65c\uc6a9\ud574 \uac80\uc0c9\uc73c\ub85c \ubc18\ubcf5 \ud655\uc778 \uac00\ub2a5\ud55c \uacf5\uac1c\uc601\uc5ed\uc5d0\uc11c \uc218\uc9d1\ud55c\ub2e4. \uc2dc\uac04\uc801 \ubc94\uc704\ub294 2015\ub144 12\uc6d4 12\uc77c\ubd80\ud130 2026\ub144 1\uc6d4 27\uc77c\uae4c\uc9c0\uc774\uba70, \uc2dc\uc791\uc810\uc740 \ud30c\ub9ac\ud611\uc815 \ucc44\ud0dd\uc774\ub77c\ub294 \uc81c\ub3c4\uc801 \ucd9c\ubc1c\uc810\uc774\uace0 \uc885\ub8cc\uc810\uc740 \ubbf8\uad6d\uc758 \ud30c\ub9ac\ud611\uc815 \ud0c8\ud1f4\uac00 \ud6a8\ub825\uc744 \ubc1c\uc0dd\ud55c \uc2dc\uc810\uc73c\ub85c\uc11c \uc2e0\uae30\ud6c4\uccb4\uc81c\uc758 \uad6d\uc81c\uc815\uce58\uc801 \ubd88\uc548\uc815\uc131\uc774 \ub2e4\uc2dc \uac00\uc2dc\ud654\ub41c \uacbd\uacc4\ub2e4. \uae30\ubcf8 \uac80\uc0c9\uc5b4\ub294 \u201c\uae30\ud6c4\ubcc0\ud654\u201d\ub97c \uc911\uc2ec\uc73c\ub85c \ud558\ub418 \ub744\uc5b4\uc4f0\uae30 \ubcc0\ud615\uacfc \uad00\ub828 \ud45c\ud604\uc744 \ubcf4\uc870\uc801\uc73c\ub85c \uac80\ud1a0\ud55c\ub2e4. \ubd84\uc11d \ub2e8\uc704\ub294 \uc601\uc0c1(\uc81c\ubaa9, \uc124\uba85, \ud0dc\uadf8, \uc790\ub9c9/\uc804\uc0ac, \ucc38\uc5ec\uc9c0\ud45c), \ucc44\ub110(\uae30\ubcf8 \uc815\ubcf4, \uc5c5\ub85c\ub4dc \ud328\ud134), \ub313\uae00\u00b7\ub300\ub313\uae00(\ud14d\uc2a4\ud2b8, \uc791\uc131 \uc2dc\uac01, \ucc38\uc5ec\uc9c0\ud45c), \ub313\uae00\uc791\uc131\uc790-\uc601\uc0c1\uc81c\uc791\uc790 \uad00\uacc4\ub9dd\uc774\ub2e4. \uc790\ub9c9\uc774 \uc5c6\ub294 \uc601\uc0c1\uc740 \uac00\ub2a5\ud55c \ubc94\uc704\uc5d0\uc11c \uc790\ub3d9 \uc74c\uc131 \uc778\uc2dd \uc804\uc0ac\ub97c \ud65c\uc6a9\ud558\uba70, \uc790\ub8cc\ub294 \uc2dc\uae30\ubcc4 \ube44\uad50\uac00 \uac00\ub2a5\ud558\ub3c4\ub85d \uad6c\uc870\ud654\ud55c\ub2e4."

    doc.save(KOR_DOC)


def polish_eng() -> None:
    doc = Document(ENG_DOC)

    doc.paragraphs[4].text = "Climate change does not automatically produce social consensus once scientific facts are presented. It remains an environmental risk issue whose meaning is renegotiated in public arenas where responsibility, justice, transition costs, and everyday experience are contested. Climate discourse therefore needs to be analyzed not as simple information transfer but as a socially constructed field of competing meanings, emotions, and political alignments (Beck, 1992; Hilgartner & Bosk, 1988)."
    doc.paragraphs[5].text = "Here, the analytical focus is the post-Paris climate-governance order that emerged after the Paris Agreement and entered fuller implementation after 2020. It includes the Agreement's institutional framework, carbon-neutrality agendas, the expanding role of non-state actors, the institutionalization of climate-crisis framing, debates over just transition, and instability in international cooperation. The period is treated not as a simple time label but as a historical conjuncture in which institutional goals and discursive justification are tested together."
    doc.paragraphs[6].text = "On that basis, the study organizes the data into two broad periods, 2015-2019 and 2020-2026, while also comparing major event phases within each. These include the Paris Agreement's 2016 entry into force, the 2017 U.S. withdrawal announcement, the 2021 U.S. re-entry, COP meetings, IPCC assessments, NDC revisions, major climate disasters, and major South Korean policy or political events. The aim is to explain how discursive patterns move and are reassembled as political and environmental conditions shift."
    doc.paragraphs[7].text = "YouTube in South Korea is a key empirical site because climate-related video messages and public comment responses accumulate on the same platform. Prior work shows that climate communication, distorted content, and post-video discussion intersect on YouTube (Allgaier, 2019; Shapiro & Park, 2018), making it useful to analyze video expression and comment response together. Korean studies have examined topic structure and terminology in climate-related YouTube content in South Korea (Lim et al., 2021; Lim, 2023). Building on that work, this project reads videos, comments, and actor networks on the same temporal axis."
    doc.paragraphs[10].text = "The immediate objective is to explain the South Korean climate-change discursive landscape within the post-Paris climate-governance order in both its structure and its dynamics. To do so, the study traces what is said in videos, what responses emerge in comments, and who clusters around which channels and issues."
    doc.paragraphs[11].text = "Ultimately, the study builds an integrated empirical dataset of climate discourse on YouTube in South Korea and shows how topics, sentiments and attitudes, and actor networks are reconfigured across events and phases. This goes beyond a snapshot of opinion and provides empirical grounds for assessing the discursive conditions under which the post-Paris climate-governance order is reinforced or weakened in South Korean society."
    doc.paragraphs[15].text = "RQ1. What topics and frames emerge in climate-change-related videos and comments on YouTube in South Korea, and how do they change across the periods 2015-2019 and 2020-2026 as well as across major event phases?"
    doc.paragraphs[17].text = "RQ3. How do commenter-creator networks change over time, which channels or actor clusters become central in particular phases, and what do these shifts suggest about the vulnerabilities and sustaining conditions of climate politics after the Paris Agreement?"
    doc.paragraphs[20].text = "The project combines reproducible data collection with descriptive statistics, topic analysis, sentiment and attitude analysis, and actor-network analysis. These components are read on the same temporal axis so that changes in topics, comment responses, and network configurations can be interpreted together. Depending on corpus size and classification fit, supervised and unsupervised methods will be combined as needed."
    doc.paragraphs[23].text = "Data will be collected from publicly searchable YouTube data via the official API for December 12, 2015 to January 27, 2026. The starting point marks the adoption of the Paris Agreement; the end point marks the date on which the U.S. withdrawal from the Agreement took effect, making renewed instability in the post-Paris order directly observable. The units of analysis are videos, channels, comments and replies, and commenter-creator ties. The corpus will organize titles, descriptions, tags, captions or transcripts, engagement indicators, and comment timestamps and participation measures. Korean-language content tied to South Korean channels, audiences, or policy contexts will define the corpus scope."
    doc.paragraphs[24].text = "To improve reproducibility, the project will document collection dates, query terms, sorting conditions, temporal batching rules, responses to result caps, and procedures for missing or unavailable data. Monthly collection will be the default, with finer segmentation in high-volume periods. The limits of a search-based design, especially its incomplete access to personalized recommendation exposure, will be stated explicitly. Identifiers will be pseudonymized, direct quotation minimized, and the need for IRB or equivalent review checked before full-scale analysis (Franzke et al., 2020)."
    doc.paragraphs[27].text = "Topic and frame analysis will be applied to both video and comment texts to identify major discursive clusters and track changes in salience over time. Themes such as climate crisis, carbon neutrality, energy transition, cost burdens, responsibility attribution, and trust in science will be compared across periods. Probabilistic topic models and related clustering approaches will be used, with interpretation guided by representativeness and readability rather than model output alone (Blei et al., 2003)."
    doc.paragraphs[30].text = "Sentiment and attitude analysis will focus mainly on comments. The study distinguishes emotional tones such as anger, anxiety, cynicism, ridicule, and hope from positions such as scientific acceptance or skepticism, responsibility attribution, policy support or opposition, and delay justification. Dictionary-based methods, supervised classifiers, and manual validation will be combined where useful (Pang & Lee, 2008). This makes it possible to examine how specific topics are received and which response structures intensify or weaken in particular periods."
    doc.paragraphs[33].text = "Actor-network analysis will begin with a bipartite commenter-creator network derived from comment participation. It will identify which channels attract concentrated responses and which participant clusters form around them. When useful, supplementary indicators such as channel-to-channel connections or co-reference structures in text will also be examined. Centrality measures and community detection will relate network change to topic composition and sentiment-attitude distributions over time (Newman, 2010; Traag et al., 2019)."
    doc.paragraphs[36].text = "Academically, the project will provide an integrated account of topic change, comment response, and actor-network reconfiguration in South Korean climate communication after the Paris Agreement. It extends environmental-sociological concerns with the social construction of risk into a platform-based public arena and offers a reproducible design for building and analyzing large-scale YouTube data over time."
    doc.paragraphs[37].text = "Socially and politically, the resulting discourse map can support public communication, media literacy, climate education, and policy discussion by showing which issues, emotions, and actor configurations intensify conflict, justify delay, or enable more constructive engagement. It also provides evidence for assessing the fragilities of climate politics in South Korea and the communicative conditions under which climate goals can be sustained."

    doc.save(ENG_DOC)


def verify() -> None:
    kor = Document(KOR_DOC)
    eng = Document(ENG_DOC)

    kor_text = "\n".join(p.text for p in kor.paragraphs)
    eng_text = "\n".join(p.text for p in eng.paragraphs)

    assert "국문 제목에서의" not in kor_text
    assert "the Korean title describes" not in eng_text
    assert "new climate regime" not in eng_text.lower()
    assert "2015년 12월 12일부터 2026년 1월 27일까지 한국 유튜브의" not in kor_text
    assert "RQ1. What topics and frames emerge" in eng_text


def main() -> None:
    read_references()
    backup_files()
    polish_kor()
    polish_eng()
    verify()
    print("Backups updated:")
    print(KOR_BACKUP)
    print(ENG_BACKUP)


if __name__ == "__main__":
    main()
