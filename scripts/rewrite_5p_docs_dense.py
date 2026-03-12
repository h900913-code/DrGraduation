from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1] / "WhatIWrote"
DOCX = ROOT / "docx"


def backup(path: Path) -> None:
    if path.exists():
        copy2(path, path.with_name(path.stem + ".backup_before_dense_revision" + path.suffix))


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_entry(doc: Document, kind: str, text: str = "") -> None:
    p = doc.add_paragraph()
    if kind == "blank":
        return
    run = p.add_run(text)
    if kind == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.bold = True
    elif kind == "author":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run.bold = True
    elif kind in {"heading", "ref_heading"}:
        run.bold = True


def rewrite(path: Path, entries: list[tuple[str, str]]) -> None:
    doc = Document(path)
    clear_body(doc)
    for kind, text in entries:
        add_entry(doc, kind, text)
    doc.save(path)


kor_entries = [
    ("title", "신기후체제하 한국 기후변화 담론 지형의 변동 분석 : 유튜브 영상댓글행위자 네트워크를 중심으로"),
    ("author", "Jeonghyeon Park"),
    ("blank", ""),
    ("heading", "1. 연구 배경 및 문제의식"),
    (
        "body",
        "기후변화는 과학적 사실이 알려진다고 해서 자동으로 사회적 합의로 이어지는 문제가 아니다. 기후위험은 책임 귀속, 정의, 전환 비용, 생활세계의 체감 경험이 뒤얽히는 가운데 공론장에서 끊임없이 재의미화되며, 따라서 기후담론 역시 정보의 양이나 정확성만으로 설명할 수 없는 사회적 구성물로 보아야 한다(Beck, 1992; Hilgartner & Bosk, 1988). 본 연구는 이러한 문제의식에서 출발하여 한국 사회의 기후변화 담론을 정적인 분포가 아니라 시간과 사건을 거치며 재편되는 관계적 지형으로 파악하고자 한다.",
    ),
    (
        "body",
        "\uc774 \uc5f0\uad6c\uc5d0\uc11c \ub9d0\ud558\ub294 \uc2e0\uae30\ud6c4\uccb4\uc81c\ub294 \ud30c\ub9ac\ud611\uc815 \ucc44\ud0dd \uc774\ud6c4 \ud615\uc131\ub418\uace0 2020\ub144 \uc774\ud6c4 \ubcf8\uaca9\uc801\uc778 \uc774\ud589 \uad6d\uba74\uc5d0 \ub4e4\uc5b4\uc120 \uc0c8\ub85c\uc6b4 \uad6d\uc81c \uae30\ud6c4\uac70\ubc84\ub10c\uc2a4 \uc9c8\uc11c\ub97c \ub73b\ud55c\ub2e4. \uc5ec\uae30\uc5d0\ub294 \ud30c\ub9ac\ud611\uc815\uc758 \uc81c\ub3c4\uc801 \ud2c0\ubfd0 \uc544\ub2c8\ub77c \ud0c4\uc18c\uc911\ub9bd \uc758\uc81c\uc758 \ud655\uc0b0, \ube44\uad6d\uac00\ud589\uc704\uc790\uc758 \ucc38\uc5ec \ud655\ub300, \uae30\ud6c4\uc704\uae30 \ud504\ub808\uc774\ubc0d\uc758 \uc81c\ub3c4\ud654, \uc815\uc758\ub85c\uc6b4 \uc804\ud658\uc744 \ub458\ub7ec\uc2fc \ub17c\uc7c1, \uad6d\uc81c\ud611\ub825\uc758 \ubd88\uc548\uc815\uc131\uc774 \ud568\uaed8 \ud3ec\ud568\ub41c\ub2e4. \uc989 \uc2e0\uae30\ud6c4\uccb4\uc81c\ub294 \ub2e8\uc21c\ud55c \uc2dc\uae30 \ud45c\uc9c0\uac00 \uc544\ub2c8\ub77c, \uc81c\ub3c4\uc801 \ubaa9\ud45c\uc640 \ub2f4\ub860\uc801 \uc815\ub2f9\ud654\uac00 \ub3d9\uc2dc\uc5d0 \uc2dc\ud5d8\ubc1b\ub294 \uc5ed\uc0ac\uc801 \uad6d\uba74\uc774\ub2e4.",
    ),
    (
        "body",
        "이때 2015년과 2020년을 구분하는 것이 중요하다. 2015년은 파리협정 채택을 통해 새로운 기준점이 제시된 해이고, 2020년은 협정의 본격 적용과 이행 논리가 가시화되기 시작한 전환점이다. 이에 본 연구는 2015~2019년을 체제 형성과 기대의 축적 국면으로, 2020~2026년을 이행과 재조정의 국면으로 구분하고, 그 안에서 2016년 협정 발효, 2017년 미국 탈퇴 선언, 2021년 미국 재가입, COP, IPCC 보고서, NDC 갱신, 대형 기후재난, 국내 정책 및 정치 이벤트 전후를 세부 비교 단위로 설정한다. 이러한 시기 구획은 담론의 변동을 단순한 연도별 빈도 변화가 아니라 국면별 재편의 과정으로 읽기 위한 장치다.",
    ),
    (
        "body",
        "한국 유튜브는 기후 관련 영상 메시지와 공개 댓글 반응이 함께 축적되는 자료 환경이라는 점에서 본 연구의 주요 관찰 지점이 된다. 국외 연구는 유튜브에서 기후 커뮤니케이션, 허위·왜곡 정보 유통, post-video discussion이 교차한다는 점을 보여주며(Allgaier, 2019; Shapiro & Park, 2018), 영상 발화와 댓글 반응을 함께 읽는 분석의 필요성을 제기한다. 국내 연구는 한국 유튜브에서 기후 관련 콘텐츠의 주제 구조와 용어 차이를 탐색해 왔다(임연수 외, 2021; 임연수, 2023). 본 연구는 이러한 선행연구를 바탕으로 영상·댓글·행위자 네트워크를 하나의 시간축에서 함께 읽는 통합 분석 설계를 제안한다.",
    ),
    ("blank", ""),
    ("heading", "2. 연구 목적"),
    (
        "body",
        "\uc774 \uc5f0\uad6c\uc758 \uc9c1\uc811\uc801\uc778 \ubaa9\uc801\uc740 \uc2e0\uae30\ud6c4\uccb4\uc81c\ud558 \ud55c\uad6d \uae30\ud6c4\ubcc0\ud654 \ub2f4\ub860 \uc9c0\ud615\uc744 \uc815\ud0dc\uc801 \uad6c\uc870\uc640 \ub3d9\ud0dc\uc801 \ubcc0\ub3d9\uc758 \uc591\uba74\uc5d0\uc11c \uc124\uba85\ud558\ub294 \ub370 \uc788\ub2e4. \uc774\ub97c \uc704\ud574 \uc601\uc0c1\uc5d0\uc11c \ubb34\uc5c7\uc774 \ub9d0\ud574\uc9c0\ub294\uc9c0, \ub313\uae00\uc5d0\uc11c \uc5b4\ub5a4 \ubc18\uc751\uc774 \ud615\uc131\ub418\ub294\uc9c0, \uadf8\ub9ac\uace0 \ub204\uac00 \uc5b4\ub5a4 \ucc44\ub110\uacfc \uc7c1\uc810\uc744 \uc911\uc2ec\uc73c\ub85c \uacb0\uc9d1\ud558\ub294\uc9c0\ub97c \ud568\uaed8 \ucd94\uc801\ud55c\ub2e4.",
    ),
    (
        "body",
        "\uad81\uadf9\uc801\uc73c\ub85c \uc774 \uc5f0\uad6c\ub294 \ud55c\uad6d \uc720\ud29c\ube0c\uc758 \uae30\ud6c4\ub2f4\ub860\uc744 \ud558\ub098\uc758 \ud1b5\ud569\ub41c \uc790\ub8cc \uccb4\uacc4\ub85c \uad6c\ucd95\ud558\uace0, \uc0ac\uac74\uacfc \uad6d\uba74\uc744 \uac70\uce58\uba70 \ud1a0\ud53d, \uc815\uc11c\u00b7\ud0dc\ub3c4, \ud589\uc704\uc790 \ub124\ud2b8\uc6cc\ud06c\uac00 \uc5b4\ub5bb\uac8c \uc7ac\ud3b8\ub418\ub294\uc9c0 \ubc1d\ud788\uace0\uc790 \ud55c\ub2e4. \uc774\ub7ec\ud55c \ubd84\uc11d\uc740 \ud2b9\uc815 \uc2dc\uae30\uc758 \uc5ec\ub860 \uc2a4\ub0c5\uc0f7\uc744 \ub118\uc5b4, \ud55c\uad6d \uc0ac\ud68c\uc5d0\uc11c \uc2e0\uae30\ud6c4\uccb4\uc81c\uac00 \uc5b4\ub5a4 \ub2f4\ub860\uc801 \uc870\uac74\uc5d0\uc11c \uac15\ud654\ub418\uac70\ub098 \uc57d\ud654\ub418\ub294\uc9c0\ub97c \uc77d\uc5b4\ub0b4\ub294 \uacbd\ud5d8\uc801 \uadfc\uac70\ub97c \uc81c\uacf5\ud560 \uac83\uc774\ub2e4.",
    ),
    ("blank", ""),
    ("heading", "3. 연구문제"),
    ("body", "이 연구는 다음의 연구문제를 중심으로 진행된다."),
    (
        "body",
        "RQ1. 2015년 12월 12일부터 2026년 1월 27일까지 한국 유튜브의 기후변화 관련 영상과 댓글에서는 어떤 주제와 프레임이 형성되며, 2015~2019년과 2020~2026년의 대구획 및 주요 사건 국면 전후에 그 담론 지형은 어떻게 재편되는가?",
    ),
    (
        "body",
        "RQ2. 댓글 담론의 반응 구조는 어떤 정서와 태도로 조직되며, 수용과 회의, 책임 귀속, 정책 찬반, 지연 정당화와 같은 반응 양상은 토픽과 시기, 사건 국면에 따라 어떻게 변화하는가?",
    ),
    (
        "body",
        "RQ3. 댓글작성자-영상제작자 관계를 중심으로 한 행위자 네트워크는 시기별로 어떻게 형성·변동하며, 어떤 채널과 행위자 집합이 특정 논점과 반응에 반복적으로 결집하는가? 이러한 변화는 신기후체제의 현재 상태와 취약성, 그리고 그 지속 조건을 읽는 데 어떤 단서를 제공하는가?",
    ),
    ("blank", ""),
    ("heading", "4. 연구방법"),
    (
        "body",
        "본 연구는 재현가능한 자료 수집과 기초통계, 토픽 분석, 정서·태도 분석, 행위자 네트워크 분석을 결합하는 혼합적 설계를 사용한다. 분석은 단일 시점의 스냅샷이 아니라 장기 시계열과 사건 기반 비교를 함께 다루는 방식으로 진행되며, 자료 규모와 분류 적합도를 고려해 지도학습과 비지도학습을 병행할 수 있다. 또한 공개 디지털 자료를 다루는 연구라는 점에서 수집 규칙의 문서화, 해석 가능 범위의 명시, 연구윤리 준수를 방법론의 일부로 포함한다.",
    ),
    ("blank", ""),
    ("heading", "4.1 자료 수집"),
    (
        "body",
        "자료는 유튜브 공식 API를 활용해 검색으로 반복 확인 가능한 공개영역에서 수집한다. 시간적 범위는 2015년 12월 12일부터 2026년 1월 27일까지이며, 시작점은 파리협정 채택이라는 제도적 출발점이고 종료점은 미국의 파리협정 탈퇴가 효력을 발생한 시점으로서 신기후체제의 국제정치적 불안정성이 다시 가시화된 경계다. 기본 검색어는 “기후변화”를 중심으로 하되 띄어쓰기 변형과 관련 표현을 보조적으로 검토한다. 분석 단위는 영상(제목, 설명, 태그, 자막/전사, 참여지표), 채널(기본 정보, 업로드 패턴), 댓글·대댓글(텍스트, 작성 시각, 참여지표), 댓글작성자-영상제작자 관계망이다. 자막이 없는 영상은 가능한 범위에서 자동 음성 인식 전사를 활용하며, 자료는 시기별 비교가 가능하도록 구조화한다.",
    ),
    (
        "body",
        "재현가능성을 높이기 위해 수집 시점, 검색어, 정렬 기준, 기간 단위, 결과 상한 대응, 결측 처리 방식을 함께 기록한다. 월 단위를 기본 수집 단위로 설정하되 결과량이 많은 구간은 더 세분화해 누락을 줄인다. 또한 검색 기반 설계가 개인화 추천에 따라 형성되는 비검색 노출을 완전히 포착하지 못한다는 한계를 명시적으로 밝히고, 삭제·비공개·댓글 비활성화 상태는 별도 결측 상태로 관리한다. 식별자는 가명처리하고, 재식별 위험 통제와 원문 인용 최소화 원칙을 적용하며, IRB 또는 기관 심의 필요성을 사전에 점검한다(Franzke et al., 2020).",
    ),
    ("blank", ""),
    ("heading", "4.2 토픽 및 프레임 분석"),
    (
        "body",
        "토픽 및 프레임 분석은 영상 텍스트와 댓글 텍스트 모두에 적용한다. 이를 통해 기후위기, 탄소중립, 에너지 전환, 책임 귀속, 생활비 부담, 과학 신뢰, 정책 회의와 같은 주제 묶음이 각 시기와 사건 국면에서 어떻게 부상하고 약화되는지 추적한다. 분석에는 확률 기반 토픽모형과 임베딩 기반 군집화 기법을 비교 적용하고(Blei et al., 2003), 필요 시 대표 문서 검토를 병행하여 해석 가능한 수준의 프레임 묶음을 정리한다. 이 단계는 영상 담론의 이동과 댓글 반응의 이동이 어떤 지점에서 맞물리거나 어긋나는지를 보여주는 핵심 축이 된다.",
    ),
    ("blank", ""),
    ("heading", "4.3 정서 및 태도 분석"),
    (
        "body",
        "정서 및 태도 분석은 주로 댓글 자료를 중심으로 수행한다. 여기서는 분노, 불안, 냉소, 조롱, 희망과 같은 정서와, 과학 수용·회의, 책임 귀속, 정책 찬반, 지연 정당화와 같은 태도를 구분하여 파악한다. 사전 기반 접근과 지도학습 분류기, 기타 기계학습 기법을 비교 적용하되, 표본 자료에 대한 인적 검토를 통해 분류 적합도를 점검한다(Pang & Lee, 2008). 이를 통해 동일한 토픽이 시기별로 어떤 정서·태도 조합을 동반하는지, 그리고 사건 국면별로 어떤 반응 구조가 강화되거나 완화되는지를 설명할 수 있다.",
    ),
    ("blank", ""),
    ("heading", "4.4 행위자 네트워크 분석"),
    (
        "body",
        "행위자 네트워크 분석은 댓글작성자-영상제작자 이분 네트워크를 기본 구조로 삼는다. 이를 통해 어떤 채널이 반복적으로 반응을 집중시키는지, 어떤 참여자 집단이 특정 채널군과 결속하는지, 그리고 그 구성이 시기와 사건 국면에 따라 어떻게 이동하는지 파악한다. 필요할 경우 채널 간 연결, 댓글 공동참여 패턴, 영상과 댓글 속 주요 행위자의 공출현 관계를 보조 지표로 검토할 수 있다. 중심성 지표와 커뮤니티 탐지 기법을 활용해(Newman, 2010; Traag et al., 2019) 토픽과 정서·태도 분포가 어떤 행위자 결집 구조와 맞물리는지를 통합적으로 해석한다.",
    ),
    ("blank", ""),
    ("heading", "5. 기대효과 및 함의"),
    (
        "body",
        "학문적으로 이 연구는 한국 유튜브의 기후담론을 발화-반응-결집의 다층 구조로 장기 추적함으로써, 환경위험의 사회적 구성과 의미 경쟁을 디지털 공론장 자료에 근거해 실증적으로 확장한다. 특히 기존 연구에서 분절적으로 다뤄지던 내용 분석, 반응 분석, 관계 분석을 동일한 해석 틀로 연결하여, 사건 국면 전후의 담론 변동을 보다 입체적으로 설명하는 근거를 제공할 수 있다. 방법론적으로도 유튜브 공개자료를 장기 시계열 자료로 구축하고, 토픽 분석, 정서·태도 분석, 네트워크 분석을 동일 시간축에서 결합하는 재현가능한 절차를 제시한다.",
    ),
    (
        "body",
        "사회적·정책적으로는 어떤 이슈와 메시지가 어떤 정서와 반응 구조를 통해 증폭되고, 어떤 행위자 구성이 갈등을 심화하거나 숙의를 가능하게 하는지 보여줌으로써 공공 커뮤니케이션 전략 수립에 기여할 수 있다. 더 나아가 이 연구는 한국 사회에서 신기후체제가 현재 어떤 상태에 놓여 있는지, 어디에서 취약성이 드러나는지, 목표 달성을 위해 어떤 담론적·정치적·소통적 조건이 필요한지를 논의하는 경험적 기반을 제공할 것이다. 반복적으로 등장하는 주장 서사와 오해 유형을 체계화한다는 점에서 미디어 리터러시와 환경교육 자료로의 활용 가능성도 기대된다.",
    ),
    ("blank", ""),
    ("ref_heading", "참고문헌"),
]

kor_refs = [
    "Allgaier, J. (2019). Science and environmental communication on YouTube: Strategically distorted communications in online videos on climate change and climate engineering. Frontiers in Communication, 4, 36.",
    "Beck, U. (1992). Risk society: Towards a new modernity. Sage.",
    "Blei, D. M., Ng, A. Y., & Jordan, M. I. (2003). Latent Dirichlet Allocation. Journal of Machine Learning Research, 3, 993-1022.",
    "Franzke, A. S., Bechmann, A., Zimmer, M., Ess, C., & Association of Internet Researchers. (2020). Internet research: Ethical guidelines 3.0. Association of Internet Researchers.",
    "Hilgartner, S., & Bosk, C. L. (1988). The rise and fall of social problems: A public arenas model. American Journal of Sociology, 94(1), 53-78.",
    "Newman, M. E. J. (2010). Networks: An introduction. Oxford University Press.",
    "임연수. (2023). 기후변화 관련 유튜브 콘텐츠에 대한 토픽모델링. 한국융합과학회지, 12(2), 139-150.",
    "임연수, 이기영, 이진균. (2021). 유튜브에서 “기후변화”, “기후위기”, “지구온난화”는 어떻게 다뤄지는가?: 기후 문제 대응을 위한 공공커뮤니케이션 방향 모색. 광고PR실학연구, 14(3), 155-184.",
    "Pang, B., & Lee, L. (2008). Opinion mining and sentiment analysis. Foundations and Trends in Information Retrieval, 2(1-2), 1-135.",
    "Shapiro, M. A., & Park, H. W. (2018). Climate change and YouTube: Deliberation potential in post-video discussions. Environmental Communication, 12(1), 115-131.",
    "Traag, V. A., Waltman, L., & van Eck, N. J. (2019). From Louvain to Leiden: Guaranteeing well-connected communities. Scientific Reports, 9, 5233.",
]
kor_entries.extend([("body", ref) for ref in kor_refs])


eng_entries = [
    ("title", "Discursive Dynamics of Climate Change in South Korea after the Paris Agreement: A Focus on YouTube Videos, Comments, and Actor Networks"),
    ("author", "Jeonghyeon Park"),
    ("blank", ""),
    ("heading", "1. Introduction"),
    (
        "body",
        "Climate change does not automatically produce social consensus once scientific facts are presented. It remains an environmental risk issue whose meaning is renegotiated in public arenas where responsibility, justice, transition costs, and everyday experience are contested. Climate discourse therefore needs to be analyzed not as simple information transfer but as a socially constructed field of competing meanings, emotions, and political alignments (Beck, 1992; Hilgartner & Bosk, 1988).",
    ),
    (
        "body",
        "Here, the analytical focus is the post-Paris climate-governance order that emerged after the Paris Agreement and entered fuller implementation after 2020. It includes the Agreement's institutional framework, carbon-neutrality agendas, the expanding role of non-state actors, the institutionalization of climate-crisis framing, debates over just transition, and instability in international cooperation. The period is treated not as a simple time label but as a historical conjuncture in which institutional goals and discursive justification are tested together.",
    ),
    (
        "body",
        "On that basis, the study organizes the data into two broad periods, 2015-2019 and 2020-2026, while also comparing major event phases within each. These include the Paris Agreement's 2016 entry into force, the 2017 U.S. withdrawal announcement, the 2021 U.S. re-entry, COP meetings, IPCC assessments, NDC revisions, major climate disasters, and major South Korean policy or political events. The aim is to explain how discursive patterns move and are reassembled as political and environmental conditions shift.",
    ),
    (
        "body",
        "YouTube in South Korea is a key empirical site because climate-related video messages and public comment responses accumulate on the same platform. Prior work shows that climate communication, distorted content, and post-video discussion intersect on YouTube (Allgaier, 2019; Shapiro & Park, 2018), making it useful to analyze video expression and comment response together. Korean studies have examined topic structure and terminology in climate-related YouTube content in South Korea (Lim et al., 2021; Lim, 2023). Building on that work, this project reads videos, comments, and actor networks on the same temporal axis.",
    ),
    ("blank", ""),
    ("heading", "2. Research Objectives"),
    (
        "body",
        "The immediate objective is to explain the South Korean climate-change discursive landscape within the post-Paris climate-governance order in both its structure and its dynamics. To do so, the study traces what is said in videos, what responses emerge in comments, and who clusters around which channels and issues.",
    ),
    (
        "body",
        "Ultimately, the study builds an integrated empirical dataset of climate discourse on YouTube in South Korea and shows how topics, sentiments and attitudes, and actor networks are reconfigured across events and phases. This goes beyond a snapshot of opinion and provides empirical grounds for assessing the discursive conditions under which the post-Paris climate-governance order is reinforced or weakened in South Korean society.",
    ),
    ("blank", ""),
    ("heading", "3. Research Questions"),
    ("body", "To achieve these objectives, the study addresses the following questions:"),
    (
        "body",
        "RQ1. What topics and frames emerge in climate-change-related videos and comments on YouTube in South Korea, and how do they change across the periods 2015-2019 and 2020-2026 as well as across major event phases?",
    ),
    (
        "body",
        "RQ2. How do comment responses vary in sentiment and attitude, and how do acceptance, skepticism, responsibility attribution, policy support or opposition, and delay justification shift across topics and phases?",
    ),
    (
        "body",
        "RQ3. How do commenter-creator networks change over time, which channels or actor clusters become central in particular phases, and what do these shifts suggest about the vulnerabilities and sustaining conditions of climate politics after the Paris Agreement?",
    ),
    ("blank", ""),
    ("heading", "4. Research Methods"),
    (
        "body",
        "The project combines reproducible data collection with descriptive statistics, topic analysis, sentiment and attitude analysis, and actor-network analysis. These components are read on the same temporal axis so that changes in topics, comment responses, and network configurations can be interpreted together. Depending on corpus size and classification fit, supervised and unsupervised methods will be combined as needed.",
    ),
    ("blank", ""),
    ("heading", "4.1 Data Collection"),
    (
        "body",
        "Data will be collected from publicly searchable YouTube data via the official API for December 12, 2015 to January 27, 2026. The starting point marks the adoption of the Paris Agreement; the end point marks the date on which the U.S. withdrawal from the Agreement took effect, making renewed instability in the post-Paris order directly observable. The units of analysis are videos, channels, comments and replies, and commenter-creator ties. The corpus will organize titles, descriptions, tags, captions or transcripts, engagement indicators, and comment timestamps and participation measures. Korean-language content tied to South Korean channels, audiences, or policy contexts will define the corpus scope.",
    ),
    (
        "body",
        "To improve reproducibility, the project will document collection dates, query terms, sorting conditions, temporal batching rules, responses to result caps, and procedures for missing or unavailable data. Monthly collection will be the default, with finer segmentation in high-volume periods. The limits of a search-based design, especially its incomplete access to personalized recommendation exposure, will be stated explicitly. Identifiers will be pseudonymized, direct quotation minimized, and the need for IRB or equivalent review checked before full-scale analysis (Franzke et al., 2020).",
    ),
    ("blank", ""),
    ("heading", "4.2 Topic and Frame Analysis"),
    (
        "body",
        "Topic and frame analysis will be applied to both video and comment texts to identify major discursive clusters and track changes in salience over time. Themes such as climate crisis, carbon neutrality, energy transition, cost burdens, responsibility attribution, and trust in science will be compared across periods. Probabilistic topic models and related clustering approaches will be used, with interpretation guided by representativeness and readability rather than model output alone (Blei et al., 2003).",
    ),
    ("blank", ""),
    ("heading", "4.3 Sentiment and Attitude Analysis"),
    (
        "body",
        "Sentiment and attitude analysis will focus mainly on comments. The study distinguishes emotional tones such as anger, anxiety, cynicism, ridicule, and hope from positions such as scientific acceptance or skepticism, responsibility attribution, policy support or opposition, and delay justification. Dictionary-based methods, supervised classifiers, and manual validation will be combined where useful (Pang & Lee, 2008). This makes it possible to examine how specific topics are received and which response structures intensify or weaken in particular periods.",
    ),
    ("blank", ""),
    ("heading", "4.4 Actor Network Analysis"),
    (
        "body",
        "Actor-network analysis will begin with a bipartite commenter-creator network derived from comment participation. It will identify which channels attract concentrated responses and which participant clusters form around them. When useful, supplementary indicators such as channel-to-channel connections or co-reference structures in text will also be examined. Centrality measures and community detection will relate network change to topic composition and sentiment-attitude distributions over time (Newman, 2010; Traag et al., 2019).",
    ),
    ("blank", ""),
    ("heading", "5. Expected Outcomes and Implications"),
    (
        "body",
        "Academically, the project will provide an integrated account of topic change, comment response, and actor-network reconfiguration in South Korean climate communication after the Paris Agreement. It extends environmental-sociological concerns with the social construction of risk into a platform-based public arena and offers a reproducible design for building and analyzing large-scale YouTube data over time.",
    ),
    (
        "body",
        "Socially and politically, the resulting discourse map can support public communication, media literacy, climate education, and policy discussion by showing which issues, emotions, and actor configurations intensify conflict, justify delay, or enable more constructive engagement. It also provides evidence for assessing the fragilities of climate politics in South Korea and the communicative conditions under which climate goals can be sustained.",
    ),
    ("blank", ""),
    ("ref_heading", "References"),
]

eng_refs = [
    "Allgaier, J. (2019). Science and environmental communication on YouTube: Strategically distorted communications in online videos on climate change and climate engineering. Frontiers in Communication, 4, 36.",
    "Beck, U. (1992). Risk society: Towards a new modernity. Sage.",
    "Blei, D. M., Ng, A. Y., & Jordan, M. I. (2003). Latent Dirichlet Allocation. Journal of Machine Learning Research, 3, 993-1022.",
    "Franzke, A. S., Bechmann, A., Zimmer, M., Ess, C., & Association of Internet Researchers. (2020). Internet research: Ethical guidelines 3.0. Association of Internet Researchers.",
    "Hilgartner, S., & Bosk, C. L. (1988). The rise and fall of social problems: A public arenas model. American Journal of Sociology, 94(1), 53-78.",
    "Newman, M. E. J. (2010). Networks: An introduction. Oxford University Press.",
    "Lim, Y. (2023). Topic modeling of YouTube contents on climate change. Journal of the Korea Convergence Society, 12(2), 139-150.",
    "Lim, Y., Lee, G., & Lee, J. (2021). How are climate change, climate crisis, and global warming addressed on YouTube? Directions for public communication on climate issues. Journal of Practical Research in Advertising and Public Relations, 14(3), 155-184.",
    "Pang, B., & Lee, L. (2008). Opinion mining and sentiment analysis. Foundations and Trends in Information Retrieval, 2(1-2), 1-135.",
    "Shapiro, M. A., & Park, H. W. (2018). Climate change and YouTube: Deliberation potential in post-video discussions. Environmental Communication, 12(1), 115-131.",
    "Traag, V. A., Waltman, L., & van Eck, N. J. (2019). From Louvain to Leiden: Guaranteeing well-connected communities. Scientific Reports, 9, 5233.",
]
eng_entries.extend([("body", ref) for ref in eng_refs])


def main() -> None:
    targets = {
        DOCX / "Doctoral_Research_Plan_5p_KOR.docx": kor_entries,
        DOCX / "Doctoral_Research_Plan_5p_ENG.docx": eng_entries,
    }
    for path in targets:
        backup(path)
    for path, entries in targets.items():
        rewrite(path, entries)
        print(f"rewritten: {path.name}")


if __name__ == "__main__":
    main()

